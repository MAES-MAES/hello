
from docx import Document
doc = Document('/Users/rustammirzoev/Desktop/Rustam_super.docx')


name_works = [
    {
        "type": "filled",
        "data": {
            "name": "Наименование работы",
            "units": "Единица измерения",
            "count": "100",
            "price": "100",
            "total": "10000", }},


    {
        "type": "subtitle",
        "data": {
            "title": "Подзаголовок",
        }},


    {   "type": "total",
        "data": {
            "value": "2324 р."
        }}]

name_materials = [
    {
        "type": "filled",
        "data": {
            
            "name": "Наименование материала",
            "units": "Единица измерения",
            "count": "100",
            "price": "100",
            "total": "10000",
        }
    },
    {
        "type": "total",
        "data": {
            "value": "2324 р."
        }

    }
    
]
totals = "100 000 р."

#СОЗДАЕМ ФАЙЛ С ДВУМЯ ТАБЛИЦАМИ
def create_word_file_with_tables(name_works,name_materials):
    table1 = doc.tables[0]
    table2 = doc.tables[1]
    insert_info_into_table(name_works,table1)
    insert_info_into_table(name_works,table2)
    doc.save('demo21.docx')
    
#ФУНКЦИЯ ДЛЯ ДОБАВЛЕНИЯ ДАННЫХ В ТАБЛИЦУ
def insert_info_into_table(data,table1):
    for index_item,item in enumerate(data):
        row = table1.add_row()
        index_row = str(index_item+1)
        if item["type"] == "filled":
            insert_filled_row(
                row,
                index_item,
                item["data"]["name"],
                item["data"]["units"],
                item["data"]["count"],
                item["data"]["price"],
                item["data"]["total"],
            )

        elif item['type'] == 'subtitle':
            insert_subtitle_row(
                row,
                index_row,
                item['data']['title']
            )
        elif item['type'] == "total":
            insert_total_row(
                row,
                index_row,
                item['data']['value']
            )

#ФУНКЦИЯ ДЛЯ ДОБАВЛЕНИЯ ПОДЗАГОЛОВКА
def insert_subtitle_row(row,index_row,text):
    row.cells[0].text = index_row
    row.cells[1].merge(row.cells[-1])
    row.cells[1].text = text


#ФУНКЦИЯ ДЛЯ ЗАПОЛНЕНИЯ ДАННЫХ ИЗ ДАТА В  ЯЧЕЙКИ И ФОРМАТИРОВАНИЕ ТЕКСТА
def insert_filled_row(
    row,
    index_row,
    name,
    units,
    count,
    price,
    total):
    p = row.cells[0].paragraphs[0]
    run = p.add_run(index_row)
    run.bold = True
    row.cells[0].paragraphs[0].alignment = 1
    row.cells[1].text = name
    row.cells[2].text = units
    row.cells[2].paragraphs[0].bold = True
    row.cells[2].paragraphs[0].alignment = 1
    row.cells[3].text = count
    row.cells[3].paragraphs[0].bold = True
    row.cells[3].paragraphs[0].alignment = 1
    row.cells[4].text = price
    row.cells[4].paragraphs[0].bold = True
    row.cells[4].paragraphs[0].alignment = 1
    row.cells[5].text = total
    row.cells[5].paragraphs[0].bold = True
    row.cells[5].paragraphs[0].alignment = 1
    # row.cells[5].add_paragraph()
    row.cells[5].paragraphs[0].alignment = 1



#ФУНКЦИЯ ДЛЯ ДОБАВЛЕНИЯ ТОТАЛ
def insert_total_row(row,index_row,text):
    row.cells[0].text = index_row
    row.cells[1].merge(row.cells[-1])
    row.cells[1].text = text




def add_total_type(
    data,
):
    total_price = 0
    for item in data:
        if "total" in item["data"]:
            total = item["data"]["total"]
            total_price += float(total)
    data.append({
        "type": "total",
        "data": {
            "value": f"ИТОГО: {total_price} Р"
        },
    })







        
create_word_file_with_tables(name_works,name_materials)

    
    

        
        

    














